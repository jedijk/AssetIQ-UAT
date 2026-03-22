"""
Generate ReliabilityOS Architecture & Cost Documentation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_document():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('ReliabilityOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('AI-Powered Reliability Intelligence Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    
    # Document info
    info = doc.add_paragraph()
    info.add_run('Architecture & Cost Documentation').bold = True
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture',
        '   2.1 Technology Stack',
        '   2.2 Frontend Architecture',
        '   2.3 Backend Architecture', 
        '   2.4 Database Schema',
        '   2.5 Data Flow Diagram',
        '3. API Documentation',
        '   3.1 Endpoint Summary',
        '   3.2 AI-Powered Endpoints',
        '4. Cost Analysis',
        '   4.1 AI/LLM Usage Costs',
        '   4.2 Infrastructure Hosting Costs',
        '   4.3 Total Cost Projections',
        '5. Equipment Types & Failure Modes',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'ReliabilityOS is an AI-powered reliability intelligence platform designed for industrial '
        'reliability engineers. The system enables users to capture equipment failures via natural '
        'language chat, automatically structures the data according to ISO 14224 standards, and '
        'provides AI-driven risk prioritization and maintenance strategy recommendations.'
    )
    
    doc.add_paragraph('Key Features:', style='List Bullet')
    features = [
        'AI-powered threat/failure capture via chat interface',
        'Automatic risk scoring and prioritization',
        'Equipment hierarchy management (ISO 14224 compliant)',
        'Causal analysis engine for root cause investigation',
        'AI-generated maintenance strategies',
        'Comprehensive FMEA library with 215+ failure modes',
        'Multi-language support (English/Dutch)',
        'Real-time dashboards with deep-linking navigation'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. System Architecture
    doc.add_heading('2. System Architecture', level=1)
    
    # 2.1 Technology Stack
    doc.add_heading('2.1 Technology Stack', level=2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Purpose'
    
    tech_stack = [
        ('Frontend', 'React 18, TailwindCSS, Shadcn/UI', 'User interface'),
        ('State Management', 'React Context API, React Query', 'Application state'),
        ('API Client', 'Axios', 'HTTP communication'),
        ('Backend', 'FastAPI (Python 3.11)', 'REST API server'),
        ('Database', 'MongoDB Atlas (Motor async)', 'Data persistence'),
        ('AI/LLM', 'GPT-5.2 via Emergent LLM Key', 'Natural language processing'),
        ('Authentication', 'JWT Bearer Tokens', 'Security'),
        ('Deployment', 'Docker, Kubernetes', 'Container orchestration'),
    ]
    
    for layer, tech, purpose in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    doc.add_paragraph()
    
    # 2.2 Frontend Architecture
    doc.add_heading('2.2 Frontend Architecture', level=2)
    
    doc.add_paragraph('Directory Structure:', style='Heading 3')
    frontend_structure = """
/app/frontend/src/
├── App.js                    # Main application entry
├── components/
│   ├── Layout.js             # Main layout wrapper
│   ├── ChatSidebar.js        # AI chat interface
│   ├── EquipmentHierarchy.js # ISO 14224 tree view
│   ├── MaintenanceStrategiesPanel.jsx
│   ├── AIInsightsPanel.jsx
│   ├── CausalIntelligencePanel.jsx
│   ├── BackButton.jsx
│   ├── maintenance/          # Extracted components
│   │   ├── CollapsibleSection.jsx
│   │   ├── EditableItem.jsx
│   │   └── FailureModesDisplay.jsx
│   └── ui/                   # Shadcn UI components
├── contexts/
│   ├── AuthContext.js        # Authentication state
│   ├── LanguageContext.js    # EN/NL localization
│   └── UndoContext.js        # Global undo functionality
├── pages/
│   ├── DashboardPage.js      # Operational dashboard
│   ├── ReliabilityPerformancePage.js
│   ├── ThreatsPage.js        # Threat management
│   ├── CausalEnginePage.js   # Root cause analysis
│   ├── EquipmentManagerPage.js
│   ├── FailureModesPage.js   # FMEA library
│   └── ActionsPage.js        # Action tracking
└── lib/
    └── api.js                # Centralized API client
"""
    doc.add_paragraph(frontend_structure, style='No Spacing')
    
    doc.add_paragraph()
    doc.add_paragraph('Pages Overview:', style='Heading 3')
    
    pages_table = doc.add_table(rows=1, cols=3)
    pages_table.style = 'Table Grid'
    hdr = pages_table.rows[0].cells
    hdr[0].text = 'Page'
    hdr[1].text = 'Route'
    hdr[2].text = 'Description'
    
    pages = [
        ('Dashboard', '/', 'Operational overview with risk distribution'),
        ('Reliability Performance', '/reliability', 'KPIs and performance metrics'),
        ('Threats', '/threats', 'Threat capture and management'),
        ('Causal Engine', '/causal-engine', 'Root cause analysis investigations'),
        ('Equipment Manager', '/equipment-manager', 'ISO 14224 hierarchy management'),
        ('Library', '/library', 'FMEA failure modes & equipment types'),
        ('Actions', '/actions', 'Centralized action tracking'),
    ]
    
    for page, route, desc in pages:
        row = pages_table.add_row().cells
        row[0].text = page
        row[1].text = route
        row[2].text = desc
    
    doc.add_page_break()
    
    # 2.3 Backend Architecture
    doc.add_heading('2.3 Backend Architecture', level=2)
    
    doc.add_paragraph('Directory Structure:', style='Heading 3')
    backend_structure = """
/app/backend/
├── server.py                      # Main FastAPI application (~4,300 lines)
├── routes/                        # Modular API routers
│   ├── __init__.py
│   ├── deps.py                    # Shared dependencies
│   ├── auth.py                    # Authentication endpoints
│   ├── threats.py                 # Threat management
│   └── stats.py                   # Statistics & scores
├── ai_helpers.py                  # LLM integration helpers
├── ai_risk_engine.py              # Risk analysis engine
├── ai_risk_models.py              # Risk data models
├── maintenance_strategy_generator.py  # AI maintenance generation
├── maintenance_strategy_models.py
├── investigation_models.py        # Causal investigation models
├── iso14224_models.py             # Equipment type definitions
├── failure_modes.py               # FMEA library (215 modes)
└── requirements.txt
"""
    doc.add_paragraph(backend_structure, style='No Spacing')
    
    doc.add_paragraph()
    
    # 2.4 Database Schema
    doc.add_heading('2.4 Database Schema', level=2)
    
    doc.add_paragraph('MongoDB Collections (21 Total):', style='Heading 3')
    
    collections_table = doc.add_table(rows=1, cols=3)
    collections_table.style = 'Table Grid'
    hdr = collections_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Collection'
    hdr[2].text = 'Purpose'
    
    collections = [
        ('Core', 'users', 'User accounts and authentication'),
        ('Core', 'threats', 'Captured threats/failures'),
        ('Core', 'chat_messages', 'AI chat conversation history'),
        ('Core', 'actions', 'General action items'),
        ('Core', 'central_actions', 'Centralized action tracking'),
        ('Equipment', 'equipment_nodes', 'Hierarchy structure'),
        ('Equipment', 'equipment_types', 'Equipment type definitions'),
        ('Equipment', 'custom_equipment_types', 'User-defined types'),
        ('Equipment', 'unstructured_items', 'Unparsed equipment data'),
        ('Investigations', 'investigations', 'Causal analysis cases'),
        ('Investigations', 'timeline_events', 'Investigation timelines'),
        ('Investigations', 'failure_identifications', 'Identified failures'),
        ('Investigations', 'cause_nodes', 'Root cause trees'),
        ('Investigations', 'action_items', 'Investigation actions'),
        ('Investigations', 'evidence_items', 'Supporting evidence'),
        ('AI Cache', 'ai_risk_insights', 'Cached risk analysis'),
        ('AI Cache', 'ai_fault_trees', 'Generated fault trees'),
        ('AI Cache', 'ai_bow_ties', 'Bow-tie diagrams'),
        ('AI Cache', 'ai_causal_analysis', 'AI causal findings'),
        ('AI Cache', 'ai_action_optimization', 'Optimized actions'),
        ('Maintenance', 'maintenance_strategies', 'Generated strategies'),
    ]
    
    for cat, coll, purpose in collections:
        row = collections_table.add_row().cells
        row[0].text = cat
        row[1].text = coll
        row[2].text = purpose
    
    doc.add_page_break()
    
    # 2.5 Data Flow
    doc.add_heading('2.5 Data Flow Diagram', level=2)
    
    data_flow = """
USER INPUT (Text/Voice)
       │
       ▼
┌─────────────────────┐
│   AI Intent         │
│   Classification    │
│   (GPT-5.2)         │
└──────────┬──────────┘
           │
     ┌─────┴─────┐
     │           │
     ▼           ▼
┌─────────┐  ┌─────────────┐
│ DATA    │  │ THREAT      │
│ QUERY   │  │ EXTRACTION  │
└─────────┘  └──────┬──────┘
                    │
                    ▼
          ┌─────────────────┐
          │ RISK SCORING    │
          │ (Severity ×     │
          │  Probability)   │
          └────────┬────────┘
                   │
     ┌─────────────┼─────────────┐
     │             │             │
     ▼             ▼             ▼
┌─────────┐  ┌───────────┐  ┌───────────┐
│ CAUSAL  │  │ AI RISK   │  │MAINTENANCE│
│ ENGINE  │  │ ENGINE    │  │ GENERATOR │
│ (Manual │  │ (Fault    │  │ (ISO14224 │
│  RCA)   │  │  Trees,   │  │  + FMEA)  │
│         │  │  Bow-Tie) │  │           │
└────┬────┘  └─────┬─────┘  └─────┬─────┘
     │             │              │
     └─────────────┼──────────────┘
                   ▼
          ┌─────────────────┐
          │ ACTIONS CENTER  │
          │ (Tracking &     │
          │  Follow-up)     │
          └────────┬────────┘
                   │
                   ▼
          ┌─────────────────┐
          │   DASHBOARDS    │
          │ • Operational   │
          │ • Reliability   │
          └─────────────────┘
"""
    doc.add_paragraph(data_flow, style='No Spacing')
    
    doc.add_page_break()
    
    # 3. API Documentation
    doc.add_heading('3. API Documentation', level=1)
    
    doc.add_heading('3.1 Endpoint Summary', level=2)
    
    doc.add_paragraph('Total Endpoints: ~106')
    doc.add_paragraph()
    
    endpoint_table = doc.add_table(rows=1, cols=3)
    endpoint_table.style = 'Table Grid'
    hdr = endpoint_table.rows[0].cells
    hdr[0].text = 'Module'
    hdr[1].text = 'Endpoints'
    hdr[2].text = 'AI-Powered'
    
    endpoints = [
        ('Authentication', '3', 'No'),
        ('Threats (CRUD + AI)', '8', 'Yes'),
        ('Equipment Hierarchy', '15', 'No'),
        ('Failure Modes Library', '8', 'No'),
        ('Causal Engine / Investigations', '20', 'Yes'),
        ('Actions Management', '6', 'No'),
        ('AI Risk Analysis', '8', 'Yes'),
        ('AI Maintenance Strategies', '8', 'Yes'),
        ('Stats / Dashboard', '6', 'No'),
        ('Voice Transcription', '1', 'Yes'),
        ('Misc (parsing, uploads)', '~23', 'No'),
    ]
    
    for module, count, ai in endpoints:
        row = endpoint_table.add_row().cells
        row[0].text = module
        row[1].text = count
        row[2].text = ai
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 AI-Powered Endpoints', level=2)
    
    ai_endpoints = [
        ('POST /api/chat/send', 'Process chat message, extract threats'),
        ('POST /api/transcribe', 'Voice to text transcription'),
        ('POST /api/threats/{id}/analyze-risk', 'AI risk analysis'),
        ('GET /api/threats/{id}/risk-insights', 'Get cached risk insights'),
        ('POST /api/threats/{id}/generate-causes', 'Generate root causes'),
        ('POST /api/threats/{id}/generate-fault-tree', 'Generate fault tree'),
        ('POST /api/threats/{id}/generate-bow-tie', 'Generate bow-tie diagram'),
        ('POST /api/threats/{id}/optimize-actions', 'Optimize action plans'),
        ('POST /api/maintenance-strategies/generate', 'Generate single strategy'),
        ('POST /api/maintenance-strategies/generate-all', 'Generate all strategies'),
    ]
    
    ai_table = doc.add_table(rows=1, cols=2)
    ai_table.style = 'Table Grid'
    hdr = ai_table.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Description'
    
    for endpoint, desc in ai_endpoints:
        row = ai_table.add_row().cells
        row[0].text = endpoint
        row[1].text = desc
    
    doc.add_page_break()
    
    # 4. Cost Analysis
    doc.add_heading('4. Cost Analysis', level=1)
    
    doc.add_heading('4.1 AI/LLM Usage Costs', level=2)
    
    doc.add_paragraph('AI Model: GPT-5.2 via Emergent Universal Key')
    doc.add_paragraph('Pricing: ~$0.01/1K input tokens + $0.03/1K output tokens')
    doc.add_paragraph('Blended Rate: ~$0.02 per 1,000 tokens')
    doc.add_paragraph()
    
    doc.add_paragraph('Token Usage Per Feature:', style='Heading 3')
    
    token_table = doc.add_table(rows=1, cols=3)
    token_table.style = 'Table Grid'
    hdr = token_table.rows[0].cells
    hdr[0].text = 'Feature'
    hdr[1].text = 'Avg Tokens/Call'
    hdr[2].text = 'Est. Calls/User/Month'
    
    token_usage = [
        ('Threat Chat', '1,500', '20'),
        ('Risk Analysis', '2,000', '5'),
        ('Cause Generation', '1,800', '3'),
        ('Maintenance Strategy Gen', '3,000', '2'),
        ('Action Optimization', '1,500', '2'),
        ('Voice Transcription', '500', '5 (if enabled)'),
    ]
    
    for feature, tokens, calls in token_usage:
        row = token_table.add_row().cells
        row[0].text = feature
        row[1].text = tokens
        row[2].text = calls
    
    doc.add_paragraph()
    
    doc.add_paragraph('Monthly Token Consumption (per active user):', style='Heading 3')
    consumption = [
        ('Threat Chat (20 calls × 1,500)', '30,000 tokens'),
        ('Risk Analysis (5 × 2,000)', '10,000 tokens'),
        ('Cause Gen (3 × 1,800)', '5,400 tokens'),
        ('Maintenance Gen (2 × 3,000)', '6,000 tokens'),
        ('Action Optimization (2 × 1,500)', '3,000 tokens'),
        ('Total per user/month', '~54,400 tokens'),
    ]
    
    for item, value in consumption:
        doc.add_paragraph(f'{item}: {value}', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Infrastructure Hosting Costs', level=2)
    
    # Budget Tier
    doc.add_paragraph('Option 1: Budget Tier (10-50 users)', style='Heading 3')
    
    budget_table = doc.add_table(rows=1, cols=3)
    budget_table.style = 'Table Grid'
    hdr = budget_table.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Provider/Spec'
    hdr[2].text = 'Monthly Cost'
    
    budget = [
        ('Frontend', 'Vercel Free / Netlify Free', '$0'),
        ('Backend', 'Railway / Render (512MB)', '$5-7'),
        ('Database', 'MongoDB Atlas M0 (Free)', '$0'),
        ('AI (Emergent)', 'Light usage', '$15-30'),
        ('Total', '', '$20-40/month'),
    ]
    
    for service, spec, cost in budget:
        row = budget_table.add_row().cells
        row[0].text = service
        row[1].text = spec
        row[2].text = cost
    
    doc.add_paragraph()
    
    # Growth Tier
    doc.add_paragraph('Option 2: Growth Tier (50-200 users)', style='Heading 3')
    
    growth_table = doc.add_table(rows=1, cols=3)
    growth_table.style = 'Table Grid'
    hdr = growth_table.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Provider/Spec'
    hdr[2].text = 'Monthly Cost'
    
    growth = [
        ('Frontend', 'Vercel Pro', '$20'),
        ('Backend', 'Railway Pro (2GB RAM)', '$25-40'),
        ('Database', 'MongoDB Atlas M10', '$57'),
        ('AI (Emergent)', 'Moderate usage', '$80-150'),
        ('Total', '', '$180-270/month'),
    ]
    
    for service, spec, cost in growth:
        row = growth_table.add_row().cells
        row[0].text = service
        row[1].text = spec
        row[2].text = cost
    
    doc.add_paragraph()
    
    # Enterprise Tier
    doc.add_paragraph('Option 3: Enterprise Tier (200-500+ users)', style='Heading 3')
    
    enterprise_table = doc.add_table(rows=1, cols=3)
    enterprise_table.style = 'Table Grid'
    hdr = enterprise_table.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Provider/Spec'
    hdr[2].text = 'Monthly Cost'
    
    enterprise = [
        ('Frontend', 'Vercel Enterprise / AWS CloudFront', '$50-100'),
        ('Backend', 'AWS ECS / GCP Cloud Run', '$100-200'),
        ('Database', 'MongoDB Atlas M30', '$210'),
        ('AI (Emergent)', 'Heavy usage', '$200-400'),
        ('Extras', 'Monitoring, backups, SSL', '$50-100'),
        ('Total', '', '$600-1,000/month'),
    ]
    
    for service, spec, cost in enterprise:
        row = enterprise_table.add_row().cells
        row[0].text = service
        row[1].text = spec
        row[2].text = cost
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 Total Cost Projections', level=2)
    
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Scale'
    hdr[1].text = 'Users'
    hdr[2].text = 'Monthly'
    hdr[3].text = 'Annual'
    
    projections = [
        ('MVP/Pilot', '10-50', '$20-40', '$240-480'),
        ('Growth', '50-200', '$180-270', '$2,200-3,200'),
        ('Enterprise', '200-500+', '$600-1,000', '$7,200-12,000'),
    ]
    
    for scale, users, monthly, annual in projections:
        row = summary_table.add_row().cells
        row[0].text = scale
        row[1].text = users
        row[2].text = monthly
        row[3].text = annual
    
    doc.add_paragraph()
    
    doc.add_paragraph('Cost Breakdown for 100 Users/Year:', style='Heading 3')
    breakdown = [
        ('Frontend (Vercel Pro)', '$240'),
        ('Backend (Railway/Render)', '$300-500'),
        ('Database (Atlas M10)', '$684'),
        ('AI/LLM (GPT-5.2)', '$1,500-2,000'),
        ('TOTAL', '$2,700 - $3,400/year'),
    ]
    
    for item, cost in breakdown:
        p = doc.add_paragraph(style='List Bullet')
        if item == 'TOTAL':
            run = p.add_run(f'{item}: {cost}')
            run.bold = True
        else:
            p.add_run(f'{item}: {cost}')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Note: ~50-60% of costs are AI/LLM usage. Infrastructure is relatively inexpensive.',
        style='Intense Quote'
    )
    
    doc.add_page_break()
    
    # 5. Equipment Types & Failure Modes
    doc.add_heading('5. Equipment Types & Failure Modes', level=1)
    
    doc.add_paragraph('Total Equipment Types: 22')
    doc.add_paragraph('Total Failure Modes: 215')
    doc.add_paragraph()
    
    doc.add_heading('5.1 Equipment Types (ISO 14224 Compliant)', level=2)
    
    equip_table = doc.add_table(rows=1, cols=4)
    equip_table.style = 'Table Grid'
    hdr = equip_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Name'
    hdr[2].text = 'ISO Class'
    hdr[3].text = 'Discipline'
    
    equipment_types = [
        ('pump_centrifugal', 'Centrifugal Pump', '1.1.1', 'Mechanical'),
        ('pump_reciprocating', 'Reciprocating Pump', '1.1.2', 'Mechanical'),
        ('compressor_centrifugal', 'Centrifugal Compressor', '1.2.1', 'Mechanical'),
        ('compressor_reciprocating', 'Reciprocating Compressor', '1.2.2', 'Mechanical'),
        ('turbine_gas', 'Gas Turbine', '1.3.1', 'Mechanical'),
        ('turbine_steam', 'Steam Turbine', '1.3.2', 'Mechanical'),
        ('extruder', 'Extruder', '1.4.1', 'Mechanical'),
        ('grinder', 'Grinder', '1.5.1', 'Mechanical'),
        ('heat_exchanger', 'Heat Exchanger', '2.1.1', 'Process'),
        ('vessel_pressure', 'Pressure Vessel', '2.2.1', 'Process'),
        ('vessel_storage', 'Storage Tank', '2.2.2', 'Process'),
        ('valve_control', 'Control Valve', '3.1.1', 'Instrumentation'),
        ('valve_safety', 'Safety Valve', '3.1.2', 'Mechanical'),
        ('valve_manual', 'Manual Valve', '3.1.3', 'Mechanical'),
        ('motor_electric', 'Electric Motor', '4.1.1', 'Electrical'),
        ('transformer', 'Transformer', '4.2.1', 'Electrical'),
        ('switchgear', 'Switchgear', '4.3.1', 'Electrical'),
        ('sensor_pressure', 'Pressure Sensor', '5.1.1', 'Instrumentation'),
        ('sensor_temperature', 'Temperature Sensor', '5.1.2', 'Instrumentation'),
        ('sensor_flow', 'Flow Sensor', '5.1.3', 'Instrumentation'),
        ('plc', 'PLC Controller', '5.2.1', 'Instrumentation'),
        ('pipe', 'Piping', '6.1.1', 'Mechanical'),
    ]
    
    for id, name, iso, disc in equipment_types:
        row = equip_table.add_row().cells
        row[0].text = id
        row[1].text = name
        row[2].text = iso
        row[3].text = disc
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Grinder Failure Modes (New)', level=2)
    
    grinder_table = doc.add_table(rows=1, cols=4)
    grinder_table.style = 'Table Grid'
    hdr = grinder_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Failure Mode'
    hdr[2].text = 'RPN'
    hdr[3].text = 'Recommended Actions'
    
    grinder_modes = [
        ('201', 'Grinding Wheel Wear', '224', 'Monitor diameter, Replace worn wheels'),
        ('202', 'Grinding Wheel Breakage', '180', 'Ring test, Check RPM rating'),
        ('203', 'Spindle Bearing Failure', '180', 'Monitor vibration, Proper lubrication'),
        ('204', 'Spindle Runout', '160', 'Measure runout, Check mounting'),
        ('205', 'Wheel Imbalance', '168', 'Balance wheel, Use balancing equipment'),
        ('206', 'Motor Overload', '200', 'Reduce feed rate, Check wheel sharpness'),
        ('207', 'Coolant System Failure', '160', 'Check coolant level, Clean filters'),
        ('208', 'Thermal Damage (Burn)', '225', 'Reduce infeed, Increase coolant'),
        ('209', 'Surface Finish Deviation', '180', 'Dress wheel, Adjust feed rate'),
        ('210', 'Dimensional Inaccuracy', '160', 'Calibrate, Check thermal stability'),
        ('211', 'Wheel Glazing', '210', 'Dress regularly, Use proper grade'),
        ('212', 'Workholding Failure', '180', 'Check chuck, Verify clamping force'),
        ('213', 'Guard Interlock Failure', '120', 'Test daily, Never bypass guards'),
        ('214', 'Dresser Failure', '150', 'Inspect dresser, Replace when worn'),
        ('215', 'Dust Extraction Failure', '160', 'Check filters, Verify airflow'),
    ]
    
    for id, mode, rpn, actions in grinder_modes:
        row = grinder_table.add_row().cells
        row[0].text = id
        row[1].text = mode
        row[2].text = rpn
        row[3].text = actions
    
    doc.add_paragraph()
    
    doc.add_heading('5.3 Failure Mode Categories', level=2)
    
    categories = [
        ('Rotating', '15 modes', 'Pumps, Compressors, Turbines, Grinders'),
        ('Static', '10 modes', 'Vessels, Heat Exchangers'),
        ('Piping', '20 modes', 'Pipes, Valves'),
        ('Instrumentation', '10 modes', 'Sensors, PLCs, Controls'),
        ('Electrical', '20 modes', 'Motors, Transformers, Switchgear'),
        ('Process', '10 modes', 'Operations, Maintenance errors'),
        ('Safety', '10 modes', 'Gas leaks, Fire, Explosion risks'),
        ('Environment', '10 modes', 'Spills, Emissions, Contamination'),
        ('Extruder', '38 modes', 'Screw, Barrel, Die, Heating systems'),
        ('Quality Control', '15 modes', 'Testing, Sampling, Documentation'),
        ('Material Handling', '10 modes', 'Storage, Ordering, Contamination'),
        ('Dosing', '10 modes', 'Feed systems, Metering'),
        ('Packaging', '15 modes', 'Weighing, Labeling, Documents'),
        ('Cooling', '6 modes', 'Water systems, Temperature control'),
        ('Mechanical', '6 modes', 'Chains, Clamps, Sieves'),
    ]
    
    cat_table = doc.add_table(rows=1, cols=3)
    cat_table.style = 'Table Grid'
    hdr = cat_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Count'
    hdr[2].text = 'Equipment Types'
    
    for cat, count, equip in categories:
        row = cat_table.add_row().cells
        row[0].text = cat
        row[1].text = count
        row[2].text = equip
    
    # Footer
    doc.add_page_break()
    doc.add_heading('Document Information', level=1)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    doc.add_paragraph('Version: 1.0')
    doc.add_paragraph('Platform: ReliabilityOS')
    doc.add_paragraph('Contact: support@reliabilityos.com')
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = '/app/ReliabilityOS_Architecture_Cost_Documentation.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
