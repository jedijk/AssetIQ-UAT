"""
Generate ReliabilityOS Full Functional Specification Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_functional_spec():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('ReliabilityOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Functional Specification Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    
    doc.add_paragraph()
    
    version_info = doc.add_paragraph('Version 1.0')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Document Control
    doc.add_heading('Document Control', level=2)
    control_table = doc.add_table(rows=4, cols=2)
    control_table.style = 'Table Grid'
    control_data = [
        ('Document Title', 'ReliabilityOS Functional Specification'),
        ('Version', '1.0'),
        ('Status', 'Final'),
        ('Last Updated', datetime.now().strftime("%Y-%m-%d")),
    ]
    for i, (label, value) in enumerate(control_data):
        control_table.rows[i].cells[0].text = label
        control_table.rows[i].cells[1].text = value
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Introduction',
        '   1.1 Purpose',
        '   1.2 Scope',
        '   1.3 Definitions & Acronyms',
        '2. Product Overview',
        '   2.1 Product Vision',
        '   2.2 Target Users',
        '   2.3 Key Benefits',
        '3. User Personas',
        '4. Functional Requirements',
        '   4.1 Authentication & User Management',
        '   4.2 Dashboard & Analytics',
        '   4.3 Threat Capture System',
        '   4.4 Equipment Hierarchy Manager',
        '   4.5 FMEA Library',
        '   4.6 Causal Analysis Engine',
        '   4.7 Maintenance Strategy Generator',
        '   4.8 Actions Management',
        '   4.9 Multi-Language Support',
        '5. User Stories',
        '6. Feature Specifications',
        '7. Data Models',
        '8. API Specifications',
        '9. Business Rules',
        '10. Non-Functional Requirements',
        '11. Appendices',
    ]
    for item in toc:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'This Functional Specification Document (FSD) defines the complete functional requirements '
        'for ReliabilityOS, an AI-powered reliability intelligence platform. It serves as the '
        'authoritative reference for development, testing, and stakeholder alignment.'
    )
    
    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        'ReliabilityOS is a comprehensive platform designed for reliability engineers in industrial '
        'environments. The system covers the full lifecycle of equipment reliability management:'
    )
    scope_items = [
        'Capture and structuring of equipment failures and threats',
        'Risk assessment and prioritization using AI',
        'Equipment hierarchy management per ISO 14224 standards',
        'Root cause analysis and causal investigation',
        'AI-generated maintenance strategies',
        'Centralized action tracking and follow-up',
        'Performance dashboards and analytics',
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('1.3 Definitions & Acronyms', level=2)
    
    def_table = doc.add_table(rows=1, cols=2)
    def_table.style = 'Table Grid'
    def_table.rows[0].cells[0].text = 'Term'
    def_table.rows[0].cells[1].text = 'Definition'
    
    definitions = [
        ('FMEA', 'Failure Mode and Effects Analysis - systematic approach for identifying potential failures'),
        ('RPN', 'Risk Priority Number - Severity x Occurrence x Detectability (1-1000 scale)'),
        ('ISO 14224', 'International standard for collection and exchange of reliability data for equipment'),
        ('RCA', 'Root Cause Analysis - method of problem solving to identify root causes of faults'),
        ('Threat', 'A potential failure, risk, or hazard identified in the system'),
        ('Causal Engine', 'Module for performing structured root cause investigations'),
        ('LLM', 'Large Language Model - AI system for natural language processing'),
        ('GPT-5.2', 'OpenAI language model used for AI features'),
        ('JWT', 'JSON Web Token - secure authentication mechanism'),
        ('CRUD', 'Create, Read, Update, Delete - basic data operations'),
    ]
    
    for term, definition in definitions:
        row = def_table.add_row().cells
        row[0].text = term
        row[1].text = definition
    
    doc.add_page_break()
    
    # 2. Product Overview
    doc.add_heading('2. Product Overview', level=1)
    
    doc.add_heading('2.1 Product Vision', level=2)
    doc.add_paragraph(
        'ReliabilityOS transforms how industrial organizations manage equipment reliability. '
        'By combining natural language AI with structured reliability engineering methodologies, '
        'the platform enables engineers to capture failures conversationally while automatically '
        'applying industry-standard frameworks like ISO 14224 and FMEA.'
    )
    doc.add_paragraph(
        'The vision is to create a self-learning reliability intelligence system that not only '
        'captures and prioritizes risks but also generates actionable maintenance strategies '
        'and identifies patterns across the entire equipment fleet.'
    )
    
    doc.add_heading('2.2 Target Users', level=2)
    
    users_table = doc.add_table(rows=1, cols=3)
    users_table.style = 'Table Grid'
    users_table.rows[0].cells[0].text = 'User Type'
    users_table.rows[0].cells[1].text = 'Role'
    users_table.rows[0].cells[2].text = 'Primary Activities'
    
    users = [
        ('Reliability Engineer', 'Primary User', 'Capture threats, analyze failures, generate strategies'),
        ('Maintenance Planner', 'Secondary User', 'Review strategies, plan maintenance activities'),
        ('Operations Manager', 'Stakeholder', 'View dashboards, track KPIs, prioritize actions'),
        ('Safety Officer', 'Reviewer', 'Review safety-critical threats, validate risk assessments'),
        ('Asset Manager', 'Strategic User', 'Monitor equipment health, make investment decisions'),
    ]
    
    for user_type, role, activities in users:
        row = users_table.add_row().cells
        row[0].text = user_type
        row[1].text = role
        row[2].text = activities
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 Key Benefits', level=2)
    benefits = [
        ('Time Savings', 'Reduce threat documentation time by 70% through AI-assisted capture'),
        ('Consistency', 'Ensure all threats are structured per ISO 14224 standards'),
        ('Risk Visibility', 'Real-time risk scoring and prioritization across all assets'),
        ('Proactive Maintenance', 'AI-generated maintenance strategies based on failure modes'),
        ('Knowledge Capture', 'Centralized repository of failure knowledge and investigations'),
        ('Multi-language', 'Full support for English and Dutch interfaces'),
    ]
    
    for benefit, description in benefits:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{benefit}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # 3. User Personas
    doc.add_heading('3. User Personas', level=1)
    
    # Persona 1
    doc.add_heading('Persona 1: Erik - Reliability Engineer', level=2)
    doc.add_paragraph('Demographics:', style='Heading 3')
    doc.add_paragraph('Age: 35-45 | Experience: 10+ years | Location: Netherlands')
    
    doc.add_paragraph('Background:', style='Heading 3')
    doc.add_paragraph(
        'Erik works at a chemical processing plant managing reliability for rotating equipment. '
        'He has deep technical knowledge but spends too much time on administrative tasks like '
        'documenting failures and creating maintenance plans.'
    )
    
    doc.add_paragraph('Goals:', style='Heading 3')
    goals = [
        'Quickly capture and document equipment failures',
        'Prioritize which issues to address first',
        'Generate defensible maintenance strategies',
        'Track root cause investigations to completion',
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_paragraph('Pain Points:', style='Heading 3')
    pains = [
        'Manual data entry is time-consuming and error-prone',
        'Difficult to maintain consistency across documentation',
        'Lacks visibility into overall risk landscape',
        'Knowledge is siloed and hard to share',
    ]
    for pain in pains:
        doc.add_paragraph(pain, style='List Bullet')
    
    doc.add_paragraph()
    
    # Persona 2
    doc.add_heading('Persona 2: Sarah - Operations Manager', level=2)
    doc.add_paragraph('Demographics:', style='Heading 3')
    doc.add_paragraph('Age: 40-50 | Experience: 15+ years | Location: Belgium')
    
    doc.add_paragraph('Background:', style='Heading 3')
    doc.add_paragraph(
        'Sarah oversees operations for multiple production units. She needs to understand '
        'equipment risks to make informed decisions about resource allocation and shutdowns.'
    )
    
    doc.add_paragraph('Goals:', style='Heading 3')
    goals = [
        'Dashboard view of all equipment risks',
        'Understand which assets need immediate attention',
        'Track completion of corrective actions',
        'Report reliability KPIs to leadership',
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Functional Requirements
    doc.add_heading('4. Functional Requirements', level=1)
    
    # 4.1 Authentication
    doc.add_heading('4.1 Authentication & User Management', level=2)
    
    auth_table = doc.add_table(rows=1, cols=4)
    auth_table.style = 'Table Grid'
    hdr = auth_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    auth_reqs = [
        ('AUTH-001', 'Users shall register with email and password', 'P0', 'Implemented'),
        ('AUTH-002', 'Users shall login with email and password', 'P0', 'Implemented'),
        ('AUTH-003', 'System shall issue JWT tokens valid for 7 days', 'P0', 'Implemented'),
        ('AUTH-004', 'Passwords shall be hashed using bcrypt', 'P0', 'Implemented'),
        ('AUTH-005', 'Protected routes shall require valid JWT', 'P0', 'Implemented'),
        ('AUTH-006', 'Users shall be able to logout (client-side)', 'P1', 'Implemented'),
        ('AUTH-007', 'System shall support password reset via email', 'P2', 'Planned'),
        ('AUTH-008', 'System shall support role-based access control', 'P2', 'Planned'),
    ]
    
    for id, req, priority, status in auth_reqs:
        row = auth_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.2 Dashboard
    doc.add_heading('4.2 Dashboard & Analytics', level=2)
    
    dash_table = doc.add_table(rows=1, cols=4)
    dash_table.style = 'Table Grid'
    hdr = dash_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    dash_reqs = [
        ('DASH-001', 'Display total threat count with breakdown by status', 'P0', 'Implemented'),
        ('DASH-002', 'Show risk distribution chart (Critical/High/Medium/Low)', 'P0', 'Implemented'),
        ('DASH-003', 'Display top 5 highest-risk threats', 'P0', 'Implemented'),
        ('DASH-004', 'Show threats by equipment type distribution', 'P1', 'Implemented'),
        ('DASH-005', 'Display threats by priority pie chart', 'P1', 'Implemented'),
        ('DASH-006', 'Implement deep-linking from dashboard cards to filtered views', 'P1', 'Implemented'),
        ('DASH-007', 'Show Reliability Performance dashboard with KPIs', 'P1', 'Implemented'),
        ('DASH-008', 'Display reliability snowflake/radar visualization', 'P2', 'Implemented'),
        ('DASH-009', 'Support drill-down navigation with back button', 'P1', 'Implemented'),
        ('DASH-010', 'Show action completion statistics', 'P1', 'Implemented'),
    ]
    
    for id, req, priority, status in dash_reqs:
        row = dash_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.3 Threat Capture
    doc.add_heading('4.3 Threat Capture System', level=2)
    
    threat_table = doc.add_table(rows=1, cols=4)
    threat_table.style = 'Table Grid'
    hdr = threat_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    threat_reqs = [
        ('THREAT-001', 'Users shall capture threats via natural language chat', 'P0', 'Implemented'),
        ('THREAT-002', 'AI shall extract asset, failure mode, and context from text', 'P0', 'Implemented'),
        ('THREAT-003', 'System shall auto-calculate risk score (severity x probability)', 'P0', 'Implemented'),
        ('THREAT-004', 'Threats shall be ranked relative to other threats', 'P0', 'Implemented'),
        ('THREAT-005', 'Users shall view, edit, and delete threats', 'P0', 'Implemented'),
        ('THREAT-006', 'System shall match threats to equipment hierarchy', 'P1', 'Implemented'),
        ('THREAT-007', 'System shall suggest relevant failure modes from FMEA library', 'P1', 'Implemented'),
        ('THREAT-008', 'Users shall attach images to threat reports', 'P2', 'Implemented'),
        ('THREAT-009', 'Users shall capture threats via voice input', 'P2', 'Implemented'),
        ('THREAT-010', 'AI shall analyze images for damage detection', 'P3', 'Planned'),
        ('THREAT-011', 'System shall support threat status workflow', 'P1', 'Implemented'),
        ('THREAT-012', 'Users shall filter and search threats', 'P1', 'Implemented'),
    ]
    
    for id, req, priority, status in threat_reqs:
        row = threat_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.4 Equipment Hierarchy
    doc.add_heading('4.4 Equipment Hierarchy Manager', level=2)
    
    equip_table = doc.add_table(rows=1, cols=4)
    equip_table.style = 'Table Grid'
    hdr = equip_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    equip_reqs = [
        ('EQUIP-001', 'Support ISO 14224 6-level hierarchy', 'P0', 'Implemented'),
        ('EQUIP-002', 'Users shall create/edit/delete equipment nodes', 'P0', 'Implemented'),
        ('EQUIP-003', 'Support drag-and-drop node reorganization', 'P1', 'Implemented'),
        ('EQUIP-004', 'Assign equipment types from predefined list', 'P0', 'Implemented'),
        ('EQUIP-005', 'Support custom equipment type creation', 'P1', 'Implemented'),
        ('EQUIP-006', 'Assign criticality levels to equipment', 'P1', 'Implemented'),
        ('EQUIP-007', 'Multi-field search (name, tag, description, type)', 'P1', 'Implemented'),
        ('EQUIP-008', 'Auto-expand parent nodes on search match', 'P1', 'Implemented'),
        ('EQUIP-009', 'Support discipline assignment (Mech/Elec/Inst/Process)', 'P1', 'Implemented'),
        ('EQUIP-010', 'Parse and import equipment lists from text', 'P2', 'Implemented'),
        ('EQUIP-011', 'Export hierarchy to PDF/Excel', 'P3', 'Planned'),
        ('EQUIP-012', 'Bulk criticality assignment', 'P3', 'Planned'),
    ]
    
    for id, req, priority, status in equip_reqs:
        row = equip_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.5 FMEA Library
    doc.add_heading('4.5 FMEA Library', level=2)
    
    fmea_table = doc.add_table(rows=1, cols=4)
    fmea_table.style = 'Table Grid'
    hdr = fmea_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    fmea_reqs = [
        ('FMEA-001', 'Provide pre-populated failure mode library (215+ modes)', 'P0', 'Implemented'),
        ('FMEA-002', 'Display failure modes with Severity, Occurrence, Detectability', 'P0', 'Implemented'),
        ('FMEA-003', 'Calculate and display RPN for each mode', 'P0', 'Implemented'),
        ('FMEA-004', 'Link failure modes to equipment types', 'P0', 'Implemented'),
        ('FMEA-005', 'Users shall create custom failure modes', 'P1', 'Implemented'),
        ('FMEA-006', 'Users shall edit existing failure modes', 'P1', 'Implemented'),
        ('FMEA-007', 'Filter modes by category and equipment type', 'P1', 'Implemented'),
        ('FMEA-008', 'Display recommended actions for each mode', 'P1', 'Implemented'),
        ('FMEA-009', 'Search failure modes by keyword', 'P1', 'Implemented'),
        ('FMEA-010', 'Manage equipment types (CRUD)', 'P1', 'Implemented'),
    ]
    
    for id, req, priority, status in fmea_reqs:
        row = fmea_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_page_break()
    
    # 4.6 Causal Engine
    doc.add_heading('4.6 Causal Analysis Engine', level=2)
    
    causal_table = doc.add_table(rows=1, cols=4)
    causal_table.style = 'Table Grid'
    hdr = causal_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    causal_reqs = [
        ('CAUSAL-001', 'Create investigation cases linked to threats', 'P0', 'Implemented'),
        ('CAUSAL-002', 'Build event timelines with dates and descriptions', 'P0', 'Implemented'),
        ('CAUSAL-003', 'Document failure identifications', 'P0', 'Implemented'),
        ('CAUSAL-004', 'Build cause trees with hierarchical nodes', 'P0', 'Implemented'),
        ('CAUSAL-005', 'Create action items from investigations', 'P0', 'Implemented'),
        ('CAUSAL-006', 'Support editable items with comments', 'P1', 'Implemented'),
        ('CAUSAL-007', 'Attach evidence to investigations', 'P1', 'Implemented'),
        ('CAUSAL-008', 'Track investigation status (Open/In Progress/Closed)', 'P1', 'Implemented'),
        ('CAUSAL-009', 'AI-generate potential causes from threat data', 'P2', 'Implemented'),
        ('CAUSAL-010', 'Generate fault tree diagrams', 'P2', 'Implemented'),
        ('CAUSAL-011', 'Generate bow-tie diagrams', 'P2', 'Implemented'),
        ('CAUSAL-012', 'Generate investigation report (PDF/PowerPoint)', 'P3', 'Planned'),
    ]
    
    for id, req, priority, status in causal_reqs:
        row = causal_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.7 Maintenance Strategy
    doc.add_heading('4.7 Maintenance Strategy Generator', level=2)
    
    maint_table = doc.add_table(rows=1, cols=4)
    maint_table.style = 'Table Grid'
    hdr = maint_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    maint_reqs = [
        ('MAINT-001', 'Generate maintenance strategies per equipment type', 'P0', 'Implemented'),
        ('MAINT-002', 'Include operator rounds and checks', 'P0', 'Implemented'),
        ('MAINT-003', 'Include scheduled maintenance tasks with frequencies', 'P0', 'Implemented'),
        ('MAINT-004', 'Include corrective action procedures', 'P0', 'Implemented'),
        ('MAINT-005', 'Include emergency response procedures', 'P1', 'Implemented'),
        ('MAINT-006', 'Link strategies to relevant failure modes', 'P1', 'Implemented'),
        ('MAINT-007', 'Support strategy versioning', 'P1', 'Implemented'),
        ('MAINT-008', 'Generate strategies for all equipment at once', 'P1', 'Implemented'),
        ('MAINT-009', 'Editable strategy items (tasks, rounds, actions)', 'P1', 'Implemented'),
        ('MAINT-010', 'Collapsible strategy card display', 'P2', 'Implemented'),
        ('MAINT-011', 'Click-through to failure mode details', 'P1', 'Implemented'),
        ('MAINT-012', 'Include spare parts recommendations', 'P2', 'Implemented'),
    ]
    
    for id, req, priority, status in maint_reqs:
        row = maint_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.8 Actions Management
    doc.add_heading('4.8 Actions Management', level=2)
    
    action_table = doc.add_table(rows=1, cols=4)
    action_table.style = 'Table Grid'
    hdr = action_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    action_reqs = [
        ('ACTION-001', 'Centralized action tracking across all modules', 'P0', 'Implemented'),
        ('ACTION-002', 'Actions shall have priority (Critical/High/Medium/Low)', 'P0', 'Implemented'),
        ('ACTION-003', 'Actions shall have status (Open/In Progress/Completed)', 'P0', 'Implemented'),
        ('ACTION-004', 'Assign owner to actions', 'P1', 'Implemented'),
        ('ACTION-005', 'Set due dates for actions', 'P1', 'Implemented'),
        ('ACTION-006', 'Filter and sort actions', 'P1', 'Implemented'),
        ('ACTION-007', 'Link actions to source (threat, investigation)', 'P1', 'Implemented'),
        ('ACTION-008', 'AI-optimize action plans', 'P2', 'Implemented'),
        ('ACTION-009', 'Action completion notifications', 'P3', 'Planned'),
    ]
    
    for id, req, priority, status in action_reqs:
        row = action_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_paragraph()
    
    # 4.9 Multi-Language
    doc.add_heading('4.9 Multi-Language Support', level=2)
    
    lang_table = doc.add_table(rows=1, cols=4)
    lang_table.style = 'Table Grid'
    hdr = lang_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Status'
    
    lang_reqs = [
        ('LANG-001', 'Support English (EN) language', 'P0', 'Implemented'),
        ('LANG-002', 'Support Dutch (NL) language', 'P0', 'Implemented'),
        ('LANG-003', 'Language toggle in application header', 'P0', 'Implemented'),
        ('LANG-004', 'Persist language preference', 'P1', 'Implemented'),
        ('LANG-005', 'Translate all UI text and labels', 'P1', 'Implemented'),
        ('LANG-006', 'Translate all dialog content', 'P1', 'Implemented'),
        ('LANG-007', 'Translate all form placeholders', 'P1', 'Implemented'),
        ('LANG-008', 'Support additional languages (DE, FR)', 'P3', 'Planned'),
    ]
    
    for id, req, priority, status in lang_reqs:
        row = lang_table.add_row().cells
        row[0].text = id
        row[1].text = req
        row[2].text = priority
        row[3].text = status
    
    doc.add_page_break()
    
    # 5. User Stories
    doc.add_heading('5. User Stories', level=1)
    
    stories = [
        ('US-001', 'Threat Capture', 
         'As a reliability engineer, I want to describe a failure in natural language so that the system automatically extracts and structures the information.',
         'High', 'Implemented'),
        ('US-002', 'Risk Prioritization',
         'As an operations manager, I want to see threats ranked by risk score so that I can prioritize which issues to address first.',
         'High', 'Implemented'),
        ('US-003', 'Equipment Hierarchy',
         'As a reliability engineer, I want to organize equipment in a hierarchy so that I can associate failures with specific assets.',
         'High', 'Implemented'),
        ('US-004', 'Maintenance Strategy',
         'As a maintenance planner, I want AI-generated maintenance strategies so that I can implement best practices without starting from scratch.',
         'High', 'Implemented'),
        ('US-005', 'Root Cause Analysis',
         'As a reliability engineer, I want to document investigations with timelines and cause trees so that I can track root cause analysis.',
         'High', 'Implemented'),
        ('US-006', 'Dashboard Overview',
         'As an operations manager, I want a dashboard showing risk distribution so that I can quickly understand the reliability landscape.',
         'High', 'Implemented'),
        ('US-007', 'Action Tracking',
         'As a reliability engineer, I want to track corrective actions so that I can ensure issues are resolved.',
         'High', 'Implemented'),
        ('US-008', 'Multi-Language',
         'As a Dutch-speaking engineer, I want to use the application in Dutch so that I can work more efficiently.',
         'Medium', 'Implemented'),
        ('US-009', 'Voice Input',
         'As a field engineer, I want to capture threats via voice so that I can report issues hands-free.',
         'Medium', 'Implemented'),
        ('US-010', 'FMEA Lookup',
         'As a reliability engineer, I want to search failure modes so that I can find relevant historical data.',
         'Medium', 'Implemented'),
        ('US-011', 'Equipment Search',
         'As a reliability engineer, I want to search equipment by multiple fields so that I can quickly find assets.',
         'Medium', 'Implemented'),
        ('US-012', 'Deep Linking',
         'As a user, I want to click dashboard metrics to see related details so that I can drill down into the data.',
         'Medium', 'Implemented'),
    ]
    
    story_table = doc.add_table(rows=1, cols=5)
    story_table.style = 'Table Grid'
    hdr = story_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Feature'
    hdr[2].text = 'User Story'
    hdr[3].text = 'Priority'
    hdr[4].text = 'Status'
    
    for id, feature, story, priority, status in stories:
        row = story_table.add_row().cells
        row[0].text = id
        row[1].text = feature
        row[2].text = story
        row[3].text = priority
        row[4].text = status
    
    doc.add_page_break()
    
    # 6. Feature Specifications
    doc.add_heading('6. Feature Specifications', level=1)
    
    doc.add_heading('6.1 AI Chat Threat Capture', level=2)
    
    doc.add_paragraph('Description:', style='Heading 3')
    doc.add_paragraph(
        'The AI Chat system enables users to report equipment failures, risks, and threats using '
        'natural language. The system uses GPT-5.2 to analyze the input and extract structured data.'
    )
    
    doc.add_paragraph('Input:', style='Heading 3')
    doc.add_paragraph('Free-form text description of an equipment issue, optionally with images or voice.')
    
    doc.add_paragraph('Processing:', style='Heading 3')
    processing_steps = [
        'Intent Classification: Determine if user is reporting a threat or asking a question',
        'Entity Extraction: Identify asset, location, failure mode, severity indicators',
        'Hierarchy Matching: Match extracted asset to equipment hierarchy',
        'FMEA Matching: Suggest relevant failure modes from library',
        'Risk Calculation: Compute risk score based on severity and probability',
        'Ranking: Position threat relative to other threats',
    ]
    for step in processing_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph('Output:', style='Heading 3')
    output_fields = [
        'Threat title and description',
        'Linked asset (if matched)',
        'Suggested failure modes',
        'Risk score (1-100)',
        'Rank among all threats',
        'Recommended actions',
    ]
    for field in output_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Risk Scoring Algorithm', level=2)
    
    doc.add_paragraph('Formula:', style='Heading 3')
    doc.add_paragraph('Risk Score = Severity (1-10) x Probability (1-10)')
    doc.add_paragraph('Result Range: 1-100')
    
    doc.add_paragraph('Risk Categories:', style='Heading 3')
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = 'Table Grid'
    hdr = risk_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Score Range'
    hdr[2].text = 'Color'
    
    risks = [
        ('Critical', '80-100', 'Red'),
        ('High', '50-79', 'Orange'),
        ('Medium', '25-49', 'Yellow'),
        ('Low', '1-24', 'Green'),
    ]
    
    for cat, range, color in risks:
        row = risk_table.add_row().cells
        row[0].text = cat
        row[1].text = range
        row[2].text = color
    
    doc.add_paragraph()
    
    doc.add_heading('6.3 ISO 14224 Hierarchy Levels', level=2)
    
    levels_table = doc.add_table(rows=1, cols=3)
    levels_table.style = 'Table Grid'
    hdr = levels_table.rows[0].cells
    hdr[0].text = 'Level'
    hdr[1].text = 'Name'
    hdr[2].text = 'Example'
    
    levels = [
        ('1', 'Installation', 'Offshore Platform Alpha'),
        ('2', 'Plant/Unit', 'Production Unit 1'),
        ('3', 'Section/System', 'Gas Compression System'),
        ('4', 'Equipment Unit', 'Centrifugal Compressor K-101'),
        ('5', 'Subunit', 'Driver Assembly'),
        ('6', 'Maintainable Item', 'Bearing 101-A'),
    ]
    
    for level, name, example in levels:
        row = levels_table.add_row().cells
        row[0].text = level
        row[1].text = name
        row[2].text = example
    
    doc.add_page_break()
    
    # 7. Data Models
    doc.add_heading('7. Data Models', level=1)
    
    doc.add_heading('7.1 Threat Model', level=2)
    
    threat_model = doc.add_table(rows=1, cols=4)
    threat_model.style = 'Table Grid'
    hdr = threat_model.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Required'
    hdr[3].text = 'Description'
    
    threat_fields = [
        ('id', 'String (UUID)', 'Yes', 'Unique identifier'),
        ('title', 'String', 'Yes', 'Threat title'),
        ('description', 'String', 'Yes', 'Detailed description'),
        ('asset', 'String', 'No', 'Linked equipment asset'),
        ('equipment_type', 'String', 'No', 'Type of equipment'),
        ('severity', 'Integer (1-10)', 'Yes', 'Impact severity'),
        ('probability', 'Integer (1-10)', 'Yes', 'Likelihood of occurrence'),
        ('risk_score', 'Integer (1-100)', 'Yes', 'Calculated risk score'),
        ('rank', 'Integer', 'Yes', 'Position among all threats'),
        ('status', 'Enum', 'Yes', 'new/in_review/mitigated/closed'),
        ('priority', 'Enum', 'Yes', 'critical/high/medium/low'),
        ('failure_modes', 'Array', 'No', 'Linked FMEA failure modes'),
        ('actions', 'Array', 'No', 'Recommended actions'),
        ('image_url', 'String', 'No', 'Attached image URL'),
        ('user_id', 'String', 'Yes', 'Owner user ID'),
        ('created_at', 'DateTime', 'Yes', 'Creation timestamp'),
        ('updated_at', 'DateTime', 'Yes', 'Last update timestamp'),
    ]
    
    for field, type, req, desc in threat_fields:
        row = threat_model.add_row().cells
        row[0].text = field
        row[1].text = type
        row[2].text = req
        row[3].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('7.2 Equipment Node Model', level=2)
    
    equip_model = doc.add_table(rows=1, cols=4)
    equip_model.style = 'Table Grid'
    hdr = equip_model.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Required'
    hdr[3].text = 'Description'
    
    equip_fields = [
        ('id', 'String (UUID)', 'Yes', 'Unique identifier'),
        ('name', 'String', 'Yes', 'Equipment name'),
        ('description', 'String', 'No', 'Description'),
        ('tag', 'String', 'No', 'Equipment tag number'),
        ('level', 'Enum', 'Yes', 'ISO 14224 level'),
        ('parent_id', 'String', 'No', 'Parent node ID'),
        ('equipment_type_id', 'String', 'No', 'Equipment type reference'),
        ('discipline', 'Enum', 'No', 'mechanical/electrical/instrumentation/process'),
        ('criticality', 'Object', 'No', 'Criticality assignment'),
        ('order_index', 'Integer', 'Yes', 'Display order'),
        ('user_id', 'String', 'Yes', 'Owner user ID'),
    ]
    
    for field, type, req, desc in equip_fields:
        row = equip_model.add_row().cells
        row[0].text = field
        row[1].text = type
        row[2].text = req
        row[3].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('7.3 Investigation Model', level=2)
    
    inv_model = doc.add_table(rows=1, cols=4)
    inv_model.style = 'Table Grid'
    hdr = inv_model.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Required'
    hdr[3].text = 'Description'
    
    inv_fields = [
        ('id', 'String (UUID)', 'Yes', 'Unique identifier'),
        ('case_number', 'String', 'Yes', 'Human-readable case number'),
        ('title', 'String', 'Yes', 'Investigation title'),
        ('description', 'String', 'No', 'Investigation description'),
        ('threat_id', 'String', 'No', 'Linked threat ID'),
        ('status', 'Enum', 'Yes', 'open/in_progress/closed'),
        ('timeline_events', 'Array', 'Yes', 'Event timeline'),
        ('failure_identifications', 'Array', 'Yes', 'Identified failures'),
        ('cause_nodes', 'Array', 'Yes', 'Root cause tree'),
        ('action_items', 'Array', 'Yes', 'Corrective actions'),
        ('evidence', 'Array', 'No', 'Supporting evidence'),
        ('user_id', 'String', 'Yes', 'Owner user ID'),
        ('created_at', 'DateTime', 'Yes', 'Creation timestamp'),
    ]
    
    for field, type, req, desc in inv_fields:
        row = inv_model.add_row().cells
        row[0].text = field
        row[1].text = type
        row[2].text = req
        row[3].text = desc
    
    doc.add_page_break()
    
    # 8. API Specifications
    doc.add_heading('8. API Specifications', level=1)
    
    doc.add_paragraph('Base URL: /api')
    doc.add_paragraph('Authentication: JWT Bearer Token')
    doc.add_paragraph('Content-Type: application/json')
    doc.add_paragraph()
    
    doc.add_heading('8.1 Authentication Endpoints', level=2)
    
    auth_api = doc.add_table(rows=1, cols=4)
    auth_api.style = 'Table Grid'
    hdr = auth_api.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Auth'
    hdr[3].text = 'Description'
    
    auth_endpoints = [
        ('POST', '/auth/register', 'No', 'Register new user'),
        ('POST', '/auth/login', 'No', 'Login and get JWT token'),
        ('GET', '/auth/me', 'Yes', 'Get current user info'),
    ]
    
    for method, endpoint, auth, desc in auth_endpoints:
        row = auth_api.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = auth
        row[3].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('8.2 Threat Endpoints', level=2)
    
    threat_api = doc.add_table(rows=1, cols=4)
    threat_api.style = 'Table Grid'
    hdr = threat_api.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Auth'
    hdr[3].text = 'Description'
    
    threat_endpoints = [
        ('POST', '/chat/send', 'Yes', 'Send chat message, may create threat'),
        ('GET', '/chat/history/{session}', 'Yes', 'Get chat history'),
        ('POST', '/transcribe', 'Yes', 'Transcribe voice to text'),
        ('GET', '/threats', 'Yes', 'List all threats'),
        ('GET', '/threats/top', 'Yes', 'Get top ranked threats'),
        ('GET', '/threats/{id}', 'Yes', 'Get single threat'),
        ('PUT', '/threats/{id}', 'Yes', 'Update threat'),
        ('DELETE', '/threats/{id}', 'Yes', 'Delete threat'),
    ]
    
    for method, endpoint, auth, desc in threat_endpoints:
        row = threat_api.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = auth
        row[3].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('8.3 Equipment Hierarchy Endpoints', level=2)
    
    equip_api = doc.add_table(rows=1, cols=4)
    equip_api.style = 'Table Grid'
    hdr = equip_api.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Auth'
    hdr[3].text = 'Description'
    
    equip_endpoints = [
        ('GET', '/equipment-hierarchy/nodes', 'Yes', 'Get all equipment nodes'),
        ('POST', '/equipment-hierarchy/nodes', 'Yes', 'Create equipment node'),
        ('PATCH', '/equipment-hierarchy/nodes/{id}', 'Yes', 'Update node'),
        ('DELETE', '/equipment-hierarchy/nodes/{id}', 'Yes', 'Delete node'),
        ('POST', '/equipment-hierarchy/nodes/{id}/move', 'Yes', 'Move node to new parent'),
        ('POST', '/equipment-hierarchy/nodes/{id}/reorder', 'Yes', 'Reorder node'),
        ('GET', '/equipment-hierarchy/types', 'Yes', 'Get equipment types'),
        ('POST', '/equipment-hierarchy/types', 'Yes', 'Create equipment type'),
        ('GET', '/equipment-hierarchy/stats', 'Yes', 'Get hierarchy statistics'),
    ]
    
    for method, endpoint, auth, desc in equip_endpoints:
        row = equip_api.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = auth
        row[3].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('8.4 AI Endpoints', level=2)
    
    ai_api = doc.add_table(rows=1, cols=4)
    ai_api.style = 'Table Grid'
    hdr = ai_api.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Auth'
    hdr[3].text = 'Description'
    
    ai_endpoints = [
        ('POST', '/threats/{id}/analyze-risk', 'Yes', 'Run AI risk analysis'),
        ('GET', '/threats/{id}/risk-insights', 'Yes', 'Get cached risk insights'),
        ('POST', '/threats/{id}/generate-causes', 'Yes', 'AI generate root causes'),
        ('POST', '/threats/{id}/generate-fault-tree', 'Yes', 'Generate fault tree'),
        ('POST', '/threats/{id}/generate-bow-tie', 'Yes', 'Generate bow-tie diagram'),
        ('POST', '/threats/{id}/optimize-actions', 'Yes', 'Optimize action plans'),
        ('POST', '/maintenance-strategies/generate', 'Yes', 'Generate single strategy'),
        ('POST', '/maintenance-strategies/generate-all', 'Yes', 'Generate all strategies'),
    ]
    
    for method, endpoint, auth, desc in ai_endpoints:
        row = ai_api.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = auth
        row[3].text = desc
    
    doc.add_page_break()
    
    # 9. Business Rules
    doc.add_heading('9. Business Rules', level=1)
    
    rules = [
        ('BR-001', 'Threat Ranking', 'When a new threat is created, recalculate ranks for all threats of the same user based on risk score descending'),
        ('BR-002', 'Risk Score Calculation', 'Risk score = Severity x Probability. Both values must be integers 1-10'),
        ('BR-003', 'Hierarchy Validation', 'Equipment nodes can only have children of the next ISO 14224 level'),
        ('BR-004', 'Case Number Generation', 'Investigation case numbers follow format: INV-YYYY-NNNN (year + sequential number)'),
        ('BR-005', 'Action Number Generation', 'Action numbers follow format: ACT-{case_number}-NN'),
        ('BR-006', 'Strategy Versioning', 'Maintenance strategy versions follow semantic versioning (major.minor)'),
        ('BR-007', 'RPN Calculation', 'FMEA RPN = Severity x Occurrence x Detectability. Range: 1-1000'),
        ('BR-008', 'User Isolation', 'Users can only see their own threats, equipment, and investigations'),
        ('BR-009', 'Criticality Inheritance', 'If equipment has criticality, child nodes inherit unless overridden'),
        ('BR-010', 'Soft Delete', 'Deleted items are marked inactive, not physically removed'),
    ]
    
    rules_table = doc.add_table(rows=1, cols=3)
    rules_table.style = 'Table Grid'
    hdr = rules_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Rule Name'
    hdr[2].text = 'Description'
    
    for id, name, desc in rules:
        row = rules_table.add_row().cells
        row[0].text = id
        row[1].text = name
        row[2].text = desc
    
    doc.add_page_break()
    
    # 10. Non-Functional Requirements
    doc.add_heading('10. Non-Functional Requirements', level=1)
    
    doc.add_heading('10.1 Performance', level=2)
    perf = [
        'Page load time: < 2 seconds',
        'API response time: < 500ms for CRUD operations',
        'AI response time: < 10 seconds for generation tasks',
        'Support 100 concurrent users',
        'Database queries optimized with indexes',
    ]
    for item in perf:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.2 Security', level=2)
    security = [
        'All API endpoints protected with JWT authentication',
        'Passwords hashed with bcrypt (cost factor 12)',
        'HTTPS enforced for all communications',
        'Input validation on all user-supplied data',
        'CORS configured for allowed origins only',
        'Rate limiting on authentication endpoints',
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.3 Availability', level=2)
    avail = [
        'Target uptime: 99.5%',
        'Graceful degradation if AI services unavailable',
        'Error messages displayed for service failures',
        'Automatic retry for transient failures',
    ]
    for item in avail:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.4 Usability', level=2)
    usability = [
        'Responsive design for desktop and tablet',
        'Support for English and Dutch languages',
        'Consistent UI patterns across all modules',
        'Keyboard navigation support',
        'Loading indicators for async operations',
        'Toast notifications for user feedback',
    ]
    for item in usability:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.5 Scalability', level=2)
    scale = [
        'Stateless backend for horizontal scaling',
        'MongoDB for flexible schema evolution',
        'Async processing for AI tasks',
        'CDN for static asset delivery',
    ]
    for item in scale:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 11. Appendices
    doc.add_heading('11. Appendices', level=1)
    
    doc.add_heading('Appendix A: Equipment Types', level=2)
    
    equip_types = [
        ('pump_centrifugal', 'Centrifugal Pump', 'Mechanical'),
        ('pump_reciprocating', 'Reciprocating Pump', 'Mechanical'),
        ('compressor_centrifugal', 'Centrifugal Compressor', 'Mechanical'),
        ('compressor_reciprocating', 'Reciprocating Compressor', 'Mechanical'),
        ('turbine_gas', 'Gas Turbine', 'Mechanical'),
        ('turbine_steam', 'Steam Turbine', 'Mechanical'),
        ('extruder', 'Extruder', 'Mechanical'),
        ('grinder', 'Grinder', 'Mechanical'),
        ('heat_exchanger', 'Heat Exchanger', 'Process'),
        ('vessel_pressure', 'Pressure Vessel', 'Process'),
        ('vessel_storage', 'Storage Tank', 'Process'),
        ('valve_control', 'Control Valve', 'Instrumentation'),
        ('valve_safety', 'Safety Valve', 'Mechanical'),
        ('valve_manual', 'Manual Valve', 'Mechanical'),
        ('motor_electric', 'Electric Motor', 'Electrical'),
        ('transformer', 'Transformer', 'Electrical'),
        ('switchgear', 'Switchgear', 'Electrical'),
        ('sensor_pressure', 'Pressure Sensor', 'Instrumentation'),
        ('sensor_temperature', 'Temperature Sensor', 'Instrumentation'),
        ('sensor_flow', 'Flow Sensor', 'Instrumentation'),
        ('plc', 'PLC Controller', 'Instrumentation'),
        ('pipe', 'Piping', 'Mechanical'),
    ]
    
    type_table = doc.add_table(rows=1, cols=3)
    type_table.style = 'Table Grid'
    hdr = type_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Name'
    hdr[2].text = 'Discipline'
    
    for id, name, disc in equip_types:
        row = type_table.add_row().cells
        row[0].text = id
        row[1].text = name
        row[2].text = disc
    
    doc.add_paragraph()
    
    doc.add_heading('Appendix B: Failure Mode Categories', level=2)
    
    categories = [
        ('Rotating', '15 failure modes', 'Pumps, Compressors, Turbines, Grinders'),
        ('Static', '10 failure modes', 'Vessels, Heat Exchangers'),
        ('Piping', '20 failure modes', 'Pipes, Valves'),
        ('Instrumentation', '10 failure modes', 'Sensors, PLCs, Controls'),
        ('Electrical', '20 failure modes', 'Motors, Transformers, Switchgear'),
        ('Process', '10 failure modes', 'Operations, Maintenance errors'),
        ('Safety', '10 failure modes', 'Gas leaks, Fire, Explosions'),
        ('Environment', '10 failure modes', 'Spills, Emissions'),
        ('Extruder', '38 failure modes', 'Screw, Barrel, Die systems'),
        ('Quality Control', '15 failure modes', 'Testing, Sampling'),
        ('Material Handling', '10 failure modes', 'Storage, Ordering'),
        ('Dosing', '10 failure modes', 'Feed systems'),
        ('Packaging', '15 failure modes', 'Weighing, Labeling'),
        ('Cooling', '6 failure modes', 'Water systems'),
        ('Mechanical', '6 failure modes', 'Chains, Clamps'),
    ]
    
    cat_table = doc.add_table(rows=1, cols=3)
    cat_table.style = 'Table Grid'
    hdr = cat_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Count'
    hdr[2].text = 'Equipment Types'
    
    for cat, count, equip in categories:
        row = cat_table.add_row().cells
        row[0].text = cat
        row[1].text = count
        row[2].text = equip
    
    # Footer
    doc.add_page_break()
    doc.add_heading('Document Approval', level=1)
    
    approval_table = doc.add_table(rows=4, cols=4)
    approval_table.style = 'Table Grid'
    hdr = approval_table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Name'
    hdr[2].text = 'Signature'
    hdr[3].text = 'Date'
    
    roles = ['Product Owner', 'Tech Lead', 'QA Lead']
    for i, role in enumerate(roles, 1):
        approval_table.rows[i].cells[0].text = role
    
    doc.add_paragraph()
    doc.add_paragraph(f'Document generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    
    return doc

if __name__ == '__main__':
    doc = create_functional_spec()
    output_path = '/app/ReliabilityOS_Functional_Specification.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
